import shutil
import subprocess
from pathlib import Path

from docx import Document

try:
    import pdfplumber
except ImportError:
    pdfplumber = None


SUPPORTED_INPUT_EXTENSIONS = {".docx", ".pdf", ".wpd", ".let", ".ltr", ".doc"}
TEXT_LIKE_EXTENSIONS = {".let", ".ltr"}


def _find_office_converter():
    for name in ("soffice", "libreoffice"):
        path = shutil.which(name)
        if path:
            return path
    for path in (
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ):
        if Path(path).exists():
            return path
    return None


def _write_text_to_docx(text: str, output_path: Path) -> Path:
    doc = Document()
    paragraphs = [p for p in text.splitlines() if p.strip()]
    if not paragraphs:
        doc.add_paragraph("")
    else:
        for para in paragraphs:
            doc.add_paragraph(para)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def _decode_text_file(input_path: Path) -> str:
    for encoding in ("utf-8", "utf-16", "cp1252", "latin-1"):
        try:
            return input_path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return input_path.read_text(encoding="latin-1", errors="ignore")


def _group_pdf_words_into_lines(words, y_tolerance: float = 3.0):
    lines = []
    current = []
    current_top = None

    for word in sorted(words, key=lambda w: (w["top"], w["x0"])):
        if current_top is None or abs(word["top"] - current_top) <= y_tolerance:
            current.append(word)
            current_top = word["top"] if current_top is None else current_top
        else:
            lines.append(current)
            current = [word]
            current_top = word["top"]

    if current:
        lines.append(current)

    return lines


def _reconstruct_pdf_paragraphs(page):
    words = page.extract_words(
        use_text_flow=True,
        keep_blank_chars=False,
        x_tolerance=2,
        y_tolerance=3,
    )
    if not words:
        return []

    word_lines = _group_pdf_words_into_lines(words)
    lines = []
    for line_words in word_lines:
        ordered = sorted(line_words, key=lambda w: w["x0"])
        text = " ".join(w["text"] for w in ordered).strip()
        if not text:
            continue
        lines.append(
            {
                "text": text,
                "top": min(w["top"] for w in ordered),
                "x0": min(w["x0"] for w in ordered),
            }
        )

    if not lines:
        return []

    paragraphs = []
    current = [lines[0]["text"]]
    prev = lines[0]

    for line in lines[1:]:
        vertical_gap = line["top"] - prev["top"]
        indent_shift = abs(line["x0"] - prev["x0"])
        starts_new = (
            vertical_gap > 14
            or indent_shift > 18
            or prev["text"].endswith(":")
        )

        if starts_new:
            paragraphs.append(" ".join(current))
            current = [line["text"]]
        else:
            if current[-1].endswith("-"):
                current[-1] = current[-1][:-1] + line["text"]
            else:
                current.append(line["text"])

        prev = line

    if current:
        paragraphs.append(" ".join(current))

    return paragraphs


def _convert_pdf_to_docx(input_path: Path, output_path: Path) -> Path:
    if pdfplumber is None:
        raise RuntimeError("PDF support requires pdfplumber to be installed.")

    pages = []
    with pdfplumber.open(input_path) as pdf:
        for page in pdf.pages:
            paragraphs = _reconstruct_pdf_paragraphs(page)
            if paragraphs:
                pages.append("\n\n".join(paragraphs))

    return _write_text_to_docx("\n\n".join(pages), output_path)


def _convert_via_office(input_path: Path, output_dir: Path) -> Path | None:
    converter = _find_office_converter()
    if not converter:
        return None

    output_dir.mkdir(parents=True, exist_ok=True)
    result = subprocess.run(
        [
            converter,
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            str(output_dir),
            str(input_path),
        ],
        check=True,
        capture_output=True,
        text=True,
    )

    converted = output_dir / f"{input_path.stem}.docx"
    if converted.exists():
        return converted

    stdout = (result.stdout or "").strip()
    stderr = (result.stderr or "").strip()
    detail = stderr or stdout or "LibreOffice produced no output file."
    raise RuntimeError(
        f"LibreOffice could not convert {input_path.name}. {detail}"
    )


def convert_input_to_docx(input_path: Path, converted_dir: Path) -> Path:
    suffix = input_path.suffix.lower()
    output_path = converted_dir / f"{input_path.stem}.docx"

    if suffix == ".docx":
        return input_path

    if suffix == ".pdf":
        try:
            office_converted = _convert_via_office(input_path, converted_dir)
            if office_converted is not None:
                return office_converted
        except Exception:
            pass
        return _convert_pdf_to_docx(input_path, output_path)

    office_converted = _convert_via_office(input_path, converted_dir)
    if office_converted is not None:
        return office_converted

    if suffix in TEXT_LIKE_EXTENSIONS:
        return _write_text_to_docx(_decode_text_file(input_path), output_path)

    raise RuntimeError(
        f"Could not convert {input_path.name}. "
        "Install LibreOffice for .wpd/.doc conversion, or provide DOCX input."
    )


def discover_and_convert_inputs(input_dir: Path, converted_dir: Path):
    source_files = sorted(
        fp for fp in input_dir.iterdir() if fp.is_file() and fp.suffix.lower() in SUPPORTED_INPUT_EXTENSIONS
    )
    converted = []
    skipped = []
    for fp in source_files:
        try:
            docx_path = convert_input_to_docx(fp, converted_dir)
            converted.append({"source_file": fp.name, "docx_path": docx_path})
        except Exception as exc:
            skipped.append({"source_file": fp.name, "reason": str(exc)})
    return converted, skipped
