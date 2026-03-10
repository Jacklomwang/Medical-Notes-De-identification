import csv
from datetime import datetime
from pathlib import Path

from docx import Document
from tqdm import tqdm

from clinical_filter import filter_clinical_paragraphs
from clinical_treatment_redactor import (
    load_treatment_pipeline,
    protect_treatment_entities,
    restore_protected_treatments,
)
from input_converter import discover_and_convert_inputs
from ner_redactor import load_ner_pipeline, redact_with_ner
from regex_rules import redact_with_regex
from text_normalizer import normalize_hard_line_breaks


ROOT = Path(__file__).resolve().parents[1]
INPUT_DIR = ROOT / "data" / "input"
INPUT_CSV = INPUT_DIR / "letters.csv"
CONVERTED_DIR = INPUT_DIR / "_converted_docx"
OUTPUT_DIR = ROOT / "data" / "output"
OUTPUT_CSV = OUTPUT_DIR / "letter_redacted.csv"


def rebuild_letters_csv(input_records):
    rows = []
    for idx, record in enumerate(input_records, start=1):
        doc = Document(record["docx_path"])
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        kept, removed = filter_clinical_paragraphs(paragraphs)
        rows.append(
            {
                "doc_id": idx,
                "source_file": record["source_file"],
                "removed_paragraphs": len(removed),
                "text": "\n".join(kept),
            }
        )

    with INPUT_CSV.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(
            f, fieldnames=["doc_id", "source_file", "removed_paragraphs", "text"]
        )
        writer.writeheader()
        writer.writerows(rows)

    print(f"Wrote {INPUT_CSV} with {len(rows)} rows")
    return rows


def run_redaction(rows):
    treatment_nlp = load_treatment_pipeline()
    nlp = load_ner_pipeline(mode="general")
    output_rows = []

    for row in tqdm(rows, desc="Redacting"):
        text = row["text"]
        text, protected_map = protect_treatment_entities(text, treatment_nlp)
        text = redact_with_regex(text)
        text = redact_with_ner(text, nlp)
        text = restore_protected_treatments(text, protected_map)
        if Path(row["source_file"]).suffix.lower() == ".pdf":
            text = normalize_hard_line_breaks(text)
        output_rows.append({"doc_id": row["doc_id"], "redacted_text": text})

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    with OUTPUT_CSV.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=["doc_id", "redacted_text"])
        writer.writeheader()
        writer.writerows(output_rows)

    print(f"Wrote {OUTPUT_CSV} with {len(output_rows)} rows")
    return output_rows


def export_separate_docx(redacted_rows, input_records):
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    separate_dir = OUTPUT_DIR / f"separate_{ts}"
    separate_dir.mkdir(parents=True, exist_ok=True)

    source_by_id = {idx: Path(record["source_file"]) for idx, record in enumerate(input_records, start=1)}

    for row in redacted_rows:
        doc_id = int(row["doc_id"])
        source_file = source_by_id[doc_id]
        output_path = separate_dir / f"{source_file.stem}_redacted.docx"

        doc = Document()
        for para in row["redacted_text"].split("\n"):
            doc.add_paragraph(para)
        doc.save(output_path)

    print(f"Wrote separate outputs to {separate_dir}")


def main():
    INPUT_DIR.mkdir(parents=True, exist_ok=True)
    CONVERTED_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    input_records, skipped_records = discover_and_convert_inputs(INPUT_DIR, CONVERTED_DIR)
    if not input_records:
        raise FileNotFoundError(
            f"No supported input files found in {INPUT_DIR} "
            "(.docx, .pdf, .wpd, .let, .ltr, .doc)"
        )

    if skipped_records:
        print("Skipped input files:")
        for record in skipped_records:
            print(f" - {record['source_file']}: {record['reason']}")

    rows = rebuild_letters_csv(input_records)
    redacted_rows = run_redaction(rows)
    export_separate_docx(redacted_rows, input_records)


if __name__ == "__main__":
    main()
